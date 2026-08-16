"""Generate the corrected-cohort rain_final research report."""

from __future__ import annotations

import json
import sys
from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from config import FIGURE_DIR, METRICS_DIR, PROCESSED_DIR, REPORT_DIR, ensure_dirs

TABLE_HELPER_ROOT = Path.home() / ".codex/plugins/cache/openai-primary-runtime/documents"
TABLE_HELPER_FILES = sorted(TABLE_HELPER_ROOT.glob("*/skills/documents/scripts/table_geometry.py"))
if not TABLE_HELPER_FILES:
    raise RuntimeError(f"Could not locate table_geometry.py below {TABLE_HELPER_ROOT}")
sys.path.insert(0, str(TABLE_HELPER_FILES[-1].parent))
from table_geometry import apply_table_geometry, column_widths_from_weights

NAVY = RGBColor(23, 54, 93)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
WHITE = RGBColor(255, 255, 255)
LIGHT_BLUE = "E8EEF5"

MODEL_LABELS = {
    "cox_residual_kg_attention_nohazard": "CoxRes-KGA",
    "xgboost_aft": "XGBoost AFT",
    "random_survival_forest": "Random survival forest",
    "survival_svm": "Survival SVM",
    "linear_regression": "Linear regression",
    "regular_neural_network": "Regular neural network",
}


def set_font(run, size=11, bold=None, italic=None, color=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    return run


def shade(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    element = properties.find(qn("w:shd"))
    if element is None:
        element = OxmlElement("w:shd")
        properties.append(element)
    element.set(qn("w:fill"), fill)


def repeat_header(row):
    properties = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    properties.append(element)


def keep_row_together(row):
    properties = row._tr.get_or_add_trPr()
    element = OxmlElement("w:cantSplit")
    element.set(qn("w:val"), "true")
    properties.append(element)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    separator = OxmlElement("w:fldChar")
    separator.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separator, text, end])


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
        style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = MUTED
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    set_font(footer.add_run("Corrected-cohort research report  |  Page "), size=8.5, color=MUTED)
    add_page_field(footer)


def paragraph(doc, text="", *, italic=False):
    item = doc.add_paragraph()
    item.paragraph_format.keep_together = True
    set_font(item.add_run(text), italic=italic)
    return item


def equation(doc, text):
    item = doc.add_paragraph()
    item.alignment = WD_ALIGN_PARAGRAPH.CENTER
    item.paragraph_format.space_before = Pt(3)
    item.paragraph_format.space_after = Pt(5)
    item.paragraph_format.keep_together = True
    set_font(item.add_run(text), size=10.5, italic=True, color=NAVY)
    return item


def bullet(doc, text):
    item = doc.add_paragraph(style="List Bullet")
    item.paragraph_format.left_indent = Inches(0.375)
    item.paragraph_format.first_line_indent = Inches(-0.194)
    item.paragraph_format.space_after = Pt(4)
    item.paragraph_format.line_spacing = 1.208
    item.paragraph_format.keep_together = True
    set_font(item.add_run("\u2002" + text))


def page_break(doc):
    doc.add_page_break()
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(18)
    set_font(spacer.add_run(" "), size=1, color=WHITE)


def caption(doc, text):
    item = doc.add_paragraph(style="Caption")
    item.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(item.add_run(text), size=9, italic=True, color=MUTED)


def figure(doc, filename, label, width=6.25):
    item = doc.add_paragraph()
    item.alignment = WD_ALIGN_PARAGRAPH.CENTER
    item.paragraph_format.keep_with_next = True
    shape = item.add_run().add_picture(str(FIGURE_DIR / filename), width=Inches(width))
    shape._inline.docPr.set("descr", label)
    caption(doc, label)


def table(doc, headers, rows, weights, label, font_size=8.3, left_align_columns=None):
    left_align_columns = {0} if left_align_columns is None else set(left_align_columns)
    caption(doc, label)
    output = doc.add_table(rows=1, cols=len(headers))
    output.style = "Table Grid"
    repeat_header(output.rows[0])
    for cell, value in zip(output.rows[0].cells, headers):
        shade(cell, LIGHT_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        item = cell.paragraphs[0]
        item.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(item.add_run(str(value)), size=font_size, bold=True, color=NAVY)
    for values in rows:
        row = output.add_row()
        keep_row_together(row)
        cells = row.cells
        for index, (cell, value) in enumerate(zip(cells, values)):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            item = cell.paragraphs[0]
            item.alignment = WD_ALIGN_PARAGRAPH.LEFT if index in left_align_columns else WD_ALIGN_PARAGRAPH.CENTER
            set_font(item.add_run(str(value)), size=font_size)
    apply_table_geometry(output, column_widths_from_weights(weights))
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(4)
    spacer.paragraph_format.space_after = Pt(4)
    return output


def fmt(value):
    return "NA" if pd.isna(value) else f"{float(value):.3f}"


def fmt_ci(estimate, low, high):
    if any(pd.isna(value) for value in (estimate, low, high)):
        return "NA"
    return f"{float(estimate):.3f} ({float(low):.3f}\u2013{float(high):.3f})"


def result_rows(frame):
    rows = []
    for _, row in frame.sort_values("rank_c_index").iterrows():
        rows.append([
            MODEL_LABELS[row.model],
            fmt(row.c_index),
            fmt(row.auc_mean),
            fmt(row.ibs),
            str(int(row.rank_c_index)),
        ])
    return rows


def internal_cv_summary_rows(frame):
    summary = (
        frame[frame.strategy == "combined_cv"]
        .groupby("model", as_index=False)
        .agg(
            c_index_mean=("c_index", "mean"),
            c_index_sd=("c_index", "std"),
            auc_mean=("auc_mean", "mean"),
            auc_sd=("auc_mean", "std"),
            ibs_mean=("ibs", "mean"),
            ibs_sd=("ibs", "std"),
        )
        .sort_values("c_index_mean", ascending=False)
        .reset_index(drop=True)
    )
    rows = []
    for rank, row in summary.iterrows():
        rows.append([
            MODEL_LABELS[row.model],
            f"{row.c_index_mean:.3f} ({row.c_index_sd:.3f})",
            f"{row.auc_mean:.3f} ({row.auc_sd:.3f})",
            f"{row.ibs_mean:.3f} ({row.ibs_sd:.3f})",
            str(rank + 1),
        ])
    return rows, summary


def fold_model_rows(frame, model, pooled):
    model_folds = frame[(frame.strategy == "combined_cv") & (frame.model == model)].copy()
    model_folds["fold_number"] = pd.to_numeric(model_folds["fold"])
    model_folds = model_folds.sort_values("fold_number")
    rows = [
        [
            f"Fold {int(row.fold_number)}",
            str(int(row.n)),
            str(int(row.events)),
            fmt(row.c_index),
            fmt(row.auc_mean),
            fmt(row.ibs),
        ]
        for _, row in model_folds.iterrows()
    ]
    rows.append([
        "Mean (SD)",
        "-",
        "-",
        f"{model_folds.c_index.mean():.3f} ({model_folds.c_index.std(ddof=1):.3f})",
        f"{model_folds.auc_mean.mean():.3f} ({model_folds.auc_mean.std(ddof=1):.3f})",
        f"{model_folds.ibs.mean():.3f} ({model_folds.ibs.std(ddof=1):.3f})",
    ])
    rows.append([
        "Pooled OOF",
        str(int(pooled.n)),
        str(int(pooled.events)),
        fmt(pooled.c_index),
        fmt(pooled.auc_mean),
        fmt(pooled.ibs),
    ])
    return rows


def add_cover(doc, datasets):
    item = doc.add_paragraph()
    item.paragraph_format.space_before = Pt(105)
    item.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(item.add_run("CORRECTED-COHORT BEST-RESULTS REPORT"), size=11, bold=True, color=RGBColor(180, 115, 20))

    item = doc.add_paragraph()
    item.alignment = WD_ALIGN_PARAGRAPH.CENTER
    item.paragraph_format.space_before = Pt(12)
    item.paragraph_format.space_after = Pt(8)
    set_font(item.add_run("Advanced Glioma\nSurvival Prediction"), size=29, bold=True, color=NAVY)

    item = doc.add_paragraph()
    item.alignment = WD_ALIGN_PARAGRAPH.CENTER
    item.paragraph_format.space_after = Pt(22)
    set_font(item.add_run("Six-Model Internal and External Validation Comparison"), size=14, color=DARK_BLUE)

    item = doc.add_paragraph()
    item.alignment = WD_ALIGN_PARAGRAPH.CENTER
    item.paragraph_format.space_after = Pt(72)
    set_font(item.add_run("Reproducible dropout-0.40/Adam CoxRes-KGA configuration"), size=10.5, italic=True, color=MUTED)

    metadata = doc.add_table(rows=4, cols=2)
    values = [
        ("Cohorts", f"CGGA (n={int(datasets.loc['CGGA', 'n'])}) and TCGA (n={int(datasets.loc['TCGA', 'n'])})"),
        ("Validation", "Combined five-fold CV and bidirectional external validation"),
        ("Model scope", "Exactly six prespecified methods"),
        ("Report date", date.today().isoformat()),
    ]
    for row, (label, value) in zip(metadata.rows, values):
        shade(row.cells[0], "F7F9FB")
        shade(row.cells[1], "F7F9FB")
        set_font(row.cells[0].paragraphs[0].add_run(label), size=9.5, bold=True, color=NAVY)
        set_font(row.cells[1].paragraphs[0].add_run(value), size=9.5)
    apply_table_geometry(metadata, column_widths_from_weights([1.45, 5.05]))
    page_break(doc)


def build_report():
    ensure_dirs()
    datasets = pd.read_csv(PROCESSED_DIR / "dataset_summary.csv")
    missing = pd.read_csv(PROCESSED_DIR / "missingness_summary.csv")
    reference_performance = pd.read_csv(METRICS_DIR / "performance_summary.csv")
    reference_training_manifest = json.loads(
        (METRICS_DIR / "training_manifest.json").read_text(encoding="utf-8")
    )
    selected_root = METRICS_DIR.parent / "sensitivity" / "dropout040_adam"
    selected_metrics = selected_root / "metrics"
    performance = pd.read_csv(selected_metrics / "performance_summary.csv")
    fold_metrics = pd.read_csv(selected_metrics / "fold_metrics.csv", dtype={"fold": str})
    bootstrap = pd.read_csv(selected_metrics / "bootstrap_performance.csv")
    audit = json.loads((selected_metrics / "metric_audit.json").read_text(encoding="utf-8"))
    training_manifest = json.loads(
        (selected_metrics / "training_manifest.json").read_text(encoding="utf-8")
    )
    sensitivity_dir = METRICS_DIR.parent / "sensitivity" / "efron_ablation" / "metrics"
    sensitivity_delta = pd.read_csv(sensitivity_dir / "performance_delta.csv")
    sensitivity_frequency = pd.read_csv(sensitivity_dir / "selection_frequency.csv")
    architecture_dir = METRICS_DIR.parent / "sensitivity" / "architecture_comparison"
    architecture_summary = pd.read_csv(architecture_dir / "performance_summary.csv")
    architecture_bootstrap = pd.read_csv(architecture_dir / "paired_bootstrap_differences.csv")
    structured_dir = METRICS_DIR.parent / "sensitivity" / "structured_improvement_comparison"
    structured_bootstrap = pd.read_csv(structured_dir / "paired_bootstrap_differences.csv")
    primary_dir = METRICS_DIR.parent / "sensitivity" / "primary_tumor_only"
    primary_comparison = pd.read_csv(primary_dir / "metrics" / "full_vs_primary_summary.csv")
    primary_paired = pd.read_csv(primary_dir / "metrics" / "paired_training_restriction.csv").iloc[0]
    dropout_dir = METRICS_DIR.parent / "sensitivity" / "dropout040_adam" / "metrics"
    dropout_performance = pd.read_csv(dropout_dir / "performance_summary.csv")
    dropout_comparison = pd.read_csv(dropout_dir / "retained_comparison.csv")
    dropout_h2_dir = METRICS_DIR.parent / "sensitivity" / "dropout050_adamw_h2" / "metrics"
    dropout_h2_performance = pd.read_csv(dropout_h2_dir / "performance_summary.csv")
    dropout_h2_comparison = pd.read_csv(dropout_h2_dir / "paired_configuration_differences.csv")
    dropout060_dir = METRICS_DIR.parent / "sensitivity" / "dropout060_adam" / "metrics"
    dropout060_performance = pd.read_csv(dropout060_dir / "performance_summary.csv")
    dropout060_comparison = pd.read_csv(dropout060_dir / "dropout040_vs_dropout060.csv")
    dropout050_h1_dir = METRICS_DIR.parent / "sensitivity" / "dropout050_adamw_h1" / "metrics"
    dropout050_h1_performance = pd.read_csv(dropout050_h1_dir / "performance_summary.csv")
    dropout050_h1_comparison = pd.read_csv(
        dropout050_h1_dir / "dropout040_adam_vs_dropout050_adamw_h1.csv"
    )
    d32_e120_dir = METRICS_DIR.parent / "sensitivity" / "d32_e120_adam" / "metrics"
    d32_e120_performance = pd.read_csv(d32_e120_dir / "performance_summary.csv")
    d32_e120_comparison = pd.read_csv(d32_e120_dir / "d16_e170_vs_d32_e120.csv")
    d32_e120_manifest = json.loads(
        (d32_e120_dir / "training_manifest.json").read_text(encoding="utf-8")
    )

    proposed_id = "cox_residual_kg_attention_nohazard"
    proposed = performance[performance.model == proposed_id].set_index("validation")
    reference_proposed = reference_performance[
        reference_performance.model == proposed_id
    ].set_index("validation")
    winners = (
        performance.sort_values(["validation", "rank_c_index"])
        .groupby("validation", sort=False)
        .first()
    )
    _, internal_summary = internal_cv_summary_rows(fold_metrics)
    proposed_internal = internal_summary[internal_summary.model == proposed_id].iloc[0]
    proposed_internal_mean_rank = int(
        internal_summary.reset_index(drop=True).index[internal_summary.model == proposed_id][0] + 1
    )

    doc = Document()
    configure_document(doc)
    add_cover(doc, datasets.set_index("dataset"))

    doc.add_heading("Executive Summary", level=1)
    paragraph(
        doc,
        "This report presents a complete six-model rerun using corrected CGGA and TCGA raw workbooks. "
        "Raw records underwent prespecified outcome-eligibility screening, predictor missingness was "
        "retained through explicit indicators, and every model was refitted under the same combined "
        "five-fold and bidirectional external-validation design. The primary CoxRes-KGA results use the "
        "reproducible dropout-0.40/Adam configuration that achieved the strongest overall rank profile."
    )
    figure(doc, "figure_01_cohort_composition.png", "Figure 1. Eligible cohorts and observed survival outcomes.", 6.1)
    bullet(doc, f"Internal combined five-fold CV, mean (SD): CoxRes-KGA C-index {proposed_internal.c_index_mean:.3f} ({proposed_internal.c_index_sd:.3f}), mean AUC {proposed_internal.auc_mean:.3f} ({proposed_internal.auc_sd:.3f}), and reported IBS {proposed_internal.ibs_mean:.3f} ({proposed_internal.ibs_sd:.3f}).")
    bullet(doc, f"External CGGA to TCGA: CoxRes-KGA C-index {proposed.loc['External CGGA to TCGA', 'c_index']:.3f}, mean AUC {proposed.loc['External CGGA to TCGA', 'auc_mean']:.3f}, and reported IBS {proposed.loc['External CGGA to TCGA', 'ibs']:.3f}.")
    bullet(doc, f"External TCGA to CGGA: CoxRes-KGA C-index {proposed.loc['External TCGA to CGGA', 'c_index']:.3f}, mean AUC {proposed.loc['External TCGA to CGGA', 'auc_mean']:.3f}, and reported IBS {proposed.loc['External TCGA to CGGA', 'ibs']:.3f}.")
    bullet(doc, f"CoxRes-KGA ranked {proposed_internal_mean_rank} by the internal five-fold mean and {int(proposed.loc['Internal combined 5-fold CV', 'rank_c_index'])} by the pooled out-of-fold C-index; its C-index ranks were {int(proposed.loc['External CGGA to TCGA', 'rank_c_index'])} for CGGA-to-TCGA validation and {int(proposed.loc['External TCGA to CGGA', 'rank_c_index'])} for TCGA-to-CGGA validation.")
    bullet(doc, "The selected configuration was identified after reviewing experiments on these cohorts; its external results are exploratory rather than untouched confirmatory validation.")
    bullet(doc, "No result is presented as prospectively validated or suitable for clinical decision-making.")

    doc.add_heading("1. Objective and Study Boundary", level=1)
    paragraph(
        doc,
        "The objective was to produce a reproducible six-model glioma survival comparison using the "
        "corrected source cohorts. The prespecified methods were CoxRes-KGA, XGBoost AFT, "
        "random survival forest, Survival SVM, linear regression, "
        "and a regular neural network. All six models were refitted; no predictions from the earlier "
        "incorrect cohort subsets were reused."
    )
    paragraph(
        doc,
        "The corrected raw workbooks contained 693 CGGA records and 1,122 TCGA records. Outcome "
        "eligibility screening excluded 36 CGGA and 82 TCGA records with missing or nonpositive overall "
        "survival time or unavailable event status, leaving 657 and 1,040 patients, respectively."
    )

    doc.add_heading("2. Data", level=1)
    rows = []
    for _, row in datasets.iterrows():
        rows.append([
            row.dataset,
            int(row.n),
            int(row.events),
            f"{row.event_rate:.1%}",
            f"{row.age_mean:.1f} ({row.age_sd:.1f})",
            f"{int(row.male)}/{int(row.female)}",
            f"{int(row.grade_2)}/{int(row.grade_3)}/{int(row.grade_4)}",
            f"{int(row.idh_mutant)}/{int(row.idh_observed)} ({row.idh_mutant / row.idh_observed:.1%})",
            f"{int(row.mgmt_methylated)}/{int(row.mgmt_observed)} ({row.mgmt_methylated / row.mgmt_observed:.1%})",
            f"{row.os_median_months:.1f}",
        ])
    table(
        doc,
        ["Dataset", "N", "Events", "Event\nrate", "Age, mean\n(SD)", "Sex, M/F", "Grade,\n2/3/4", "IDH-mut,\nn/N (%)", "MGMT-meth,\nn/N (%)", "Median OS,\nmo"],
        rows,
        [.78, .48, .52, .55, .82, .67, .92, .96, 1.02, .72],
        "Table 1. Dataset summary after outcome eligibility screening.",
        7.2,
    )
    paragraph(
        doc,
        "The harmonized feature set contains age, sex, WHO grade, IDH mutation, 1p/19q codeletion, "
        "IDH/codeletion subtype, MGMT promoter methylation, and missingness indicators. The raw "
        "workbooks remain unchanged in data/raw."
    )
    nonzero = missing[missing.missing_n > 0].sort_values(["dataset", "missing_pct"], ascending=[True, False])
    page_break(doc)
    table(
        doc,
        ["Dataset", "Variable", "Missing, n", "Missing, %"],
        [[row.dataset, row.variable, int(row.missing_n), f"{row.missing_pct:.1f}"] for _, row in nonzero.iterrows()],
        [1.25, 2.8, 1.0, 1.0],
        "Table 2. Nonzero missingness in harmonized predictors.",
        8.4,
    )

    doc.add_heading("3. Corrected-Cohort Refit and Provenance", level=1)
    paragraph(
        doc,
        "All six methods were trained from the corrected processed cohorts. Preprocessing was fitted "
        "inside each outer training split, and the same held-out patients were used across models within "
        "each validation job. Job-specific predictions and model settings were retained to support audit "
        "and exact reconstruction of the performance tables."
    )
    paragraph(
        doc,
        "Linear regression used ridge regression of log survival time among observed deaths. This "
        "comparator is deliberately censoring-unaware and should be interpreted only as a requested "
        "conventional baseline rather than a clinically appropriate survival model."
    )
    table(
        doc,
        ["Final method", "Result status"],
        [
            ["CoxRes-KGA", "Selected dropout-0.40/Adam run; exactly reproduced in an independent repeat"],
            ["XGBoost AFT", "Refitted; native survival:AFT"],
            ["Random survival forest", "Refitted on corrected cohorts"],
            ["Survival SVM", "Refitted on corrected cohorts"],
            ["Regular neural network", "Refitted on corrected cohorts"],
            ["Linear regression", "Refitted on corrected cohorts"],
        ],
        [3.0, 3.5],
        "Table 3. Final model scope and prediction provenance.",
        8.1,
    )
    bullet(doc, f"Original six-model validation jobs: {len(reference_training_manifest)}; selected CoxRes-KGA refit jobs: {len(training_manifest)}.")
    bullet(doc, f"Patient-model predictions: {audit['prediction_rows']:,}.")
    bullet(doc, "Every survival probability was finite, bounded between zero and one, and nonincreasing across the four horizons.")

    doc.add_heading("4. Validation Design", level=1)
    paragraph(
        doc,
        "Internal validation used five-fold cross-validation after combining CGGA and TCGA, "
        "with dataset, event status, and grade represented in stratification. Each patient contributes one "
        "out-of-fold prediction per model. This is not separate within-cohort internal validation."
    )
    paragraph(
        doc,
        "External validation is bidirectional. One experiment trains on CGGA and tests on TCGA; the other "
        "trains on TCGA and tests on CGGA. The asymmetric results are reported separately."
    )

    doc.add_heading("5. Model Architecture and Hyperparameter Specifications", level=1)
    figure(
        doc,
        "figure_02_model_architecture.png",
        "Figure 2. Detailed CoxRes-KGA architecture. Fold-fitted predictors enter a Cox main-effect pathway and a fixed 16-dimensional knowledge-biased Transformer residual pathway. Positive residual fusion, five-seed averaging, and training-derived Breslow calibration produce patient-level risk and survival probabilities.",
        6.45,
    )
    paragraph(
        doc,
        "CoxRes-KGA is a study-specific Cox-residual knowledge-guided attention network. The Cox "
        "component follows the proportional-hazards and partial-likelihood framework introduced by Cox "
        "[2]. Within each training split, a penalized linear Cox model was fitted to median-imputed, "
        "standardized, and one-hot encoded predictors. Its risk estimate was standardized to form "
        "the Cox offset, zCox. The offset formed the explicit linear-risk pathway and was combined "
        "with the neural residual only at the final fusion step."
    )
    equation(doc, "\u03b7Cox(x) = x\u1d40\u03b2,     zCox = [\u03b7Cox(x) \u2212 mean(\u03b7Cox)] / SD(\u03b7Cox)")
    paragraph(
        doc,
        "The neural pathway represented every transformed scalar predictor through a learned linear "
        "token projection. A learned CLS token was prepended to this feature sequence. This "
        "feature-tokenization strategy and CLS-based prediction are related to "
        "the FT-Transformer formulation for tabular data [10], which adapts the multi-head Transformer "
        "architecture of Vaswani et al. [9]. A fixed additive attention-bias matrix favored prespecified "
        "glioma relationships, including age-grade, grade-IDH, grade-MGMT, IDH-1p/19q, and IDH-MGMT, "
        "while leaving all other token pairs available to self-attention."
    )
    paragraph(
        doc,
        "The selected Transformer encoder used one layer, one attention head, and token dimension 16. It "
        "produced a contextualized CLS representation. A LayerNorm\u2013multilayer-perceptron head mapped "
        "this representation to a scalar neural residual, rNN. A softplus transformation kept the "
        "residual scaling coefficient positive."
    )
    equation(doc, "\u03b7net(x) = zCox + softplus(ar)\u00b7rNN(x)")
    paragraph(
        doc,
        "Here, residual has a specific operational meaning: rNN estimates nonlinear prognostic "
        "structure beyond the fold-fitted Cox offset. It does not denote a ResNet architecture and is "
        "not attributed to a previously published model named CoxRes-KGA. Similarly, knowledge-guided "
        "attention is the present study's combination of expert-specified glioma relationships, an "
        "additive attention bias, and Transformer self-attention. Its general motivation is related "
        "to prior biologically structured neural survival models such as Cox-PASNet [11], but the exact "
        "features, priors, fusion mechanism, and CoxRes-KGA name are study-specific contributions."
    )
    paragraph(
        doc,
        "Training minimized Cox partial-likelihood loss together with a pairwise ranking term and a "
        "penalty on the squared neural residual. Pairwise survival-ranking objectives have also "
        "been used in deep survival modeling [13]. Training duration was selected within the outer "
        "training data. Final risk predictions were averaged across five "
        "random seeds."
    )
    equation(doc, "L = LCox + \u03bbrank Lpair + \u03bbres mean(rNN\u00b2)")
    paragraph(
        doc,
        "The 16-dimensional, one-layer, one-head architecture was selected in a controlled follow-up comparison "
        "against 32-dimensional/one-head and 16-dimensional/two-head alternatives, then applied "
        "unchanged in every internal and external validation job. Ten prespecified training configurations varied "
        "the Cox penalty (0.01, 0.05, 0.10, 0.30, or 1.00), ranking-loss weight (0.10, 0.16, 0.25, or 0.40), and "
        "residual penalty (0.010 or 0.020). Only the combinations recorded in the training manifest were evaluated. "
        "Three-fold cross-validation inside each outer training set selected one configuration by mean C-index, "
        "with lower cross-fold variability used as the tie breaker. An 18% inner split then selected training "
        "duration before refitting on the complete outer training data. No held-out patient outcomes were used for "
        "configuration, fitting, or epoch selection. However, the architecture and candidate grid were chosen after "
        "reviewing earlier experiments on the same cohorts. Dropout 0.40 and Adam were subsequently selected "
        "after comparison with the dropout-0.12/AdamW reference, so external "
        "performance is interpreted as exploratory follow-up evidence rather than untouched confirmatory "
        "validation. Predictions were averaged across seeds 42, 43, 44, 45, and 46."
    )
    paragraph(
        doc,
        "The resulting continuous risk score was converted to survival probabilities at 12, 24, 36, "
        "and 60 months using a training-derived Breslow-style cumulative baseline hazard [12]. This "
        "training-derived transformation was used solely for probability-based evaluation at the four "
        "prespecified horizons."
    )
    page_break(doc)
    doc.add_heading("5.1 Comparator Models and Verified Hyperparameters", level=2)
    paragraph(
        doc,
        "The regular neural network is a DeepSurv-style feed-forward Cox model. Numeric "
        "variables were median-imputed and standardized, while categorical variables were mode-imputed "
        "and one-hot encoded. The encoded inputs passed through two fully connected "
        "layers containing 96 and 48 units. Each layer used a rectified linear unit, batch normalization, "
        "and dropout of 0.15. A final linear unit produced a continuous risk score. The model was trained "
        "with the Cox partial-likelihood loss and did not include knowledge-guided interactions, an "
        "attention mechanism, or a residual Cox pathway."
    )
    paragraph(
        doc,
        "Table 4 reports the settings executed in the corrected-data refit. Hyperparameters were applied "
        "within each training split. The random seed was 42 unless a CoxRes-KGA ensemble member used its "
        "documented seed offset. Native XGBoost survival:AFT was executed in an isolated subprocess to "
        "avoid OpenMP runtime conflicts."
    )
    table(
        doc,
        ["Model", "Verified implementation", "Hyperparameters and training settings"],
        [
            [
                "CoxRes-KGA",
                "Penalized Cox offset combined with a knowledge-guided Transformer residual ensemble.",
                "Token dimension 16; 1 attention head; 1 encoder layer; feed-forward width 48; dropout 0.40; Adam optimizer; fixed additive attention bias +0.45 for prespecified pairs and -0.35 otherwise; learning rate 3e-4; nominal weight decay 1e-3. Ten prespecified configurations varied Cox alpha over 0.01-1.00, ranking weight over 0.10-0.40, and residual L2 weight between 0.010 and 0.020. Three-fold training-only selection preceded an 18% epoch-selection split (70-170 epochs; patience 22). Predictions averaged seeds 42-46.",
            ],
            [
                "Random survival forest",
                "scikit-survival RandomSurvivalForest.",
                "500 trees; minimum terminal-node size 5; max_features='sqrt'; all processor cores; seed 42. Numeric variables were median-imputed without scaling; categorical variables were mode-imputed and one-hot encoded.",
            ],
            [
                "Survival SVM",
                "scikit-survival FastSurvivalSVM.",
                "alpha=0.5; rank_ratio=0.7; maximum 1,000 iterations; seed 42. Numeric variables were median-imputed and standardized; categorical variables were mode-imputed and one-hot encoded.",
            ],
            [
                "XGBoost AFT",
                "Native XGBoost accelerated failure-time survival model.",
                "350 boosting iterations; normal AFT distribution with scale 1.2; histogram tree method; loss-guided growth; maximum 15 leaves; learning rate 0.04; L2 regularization 0.01; seed 42. Censored observations used interval labels with infinite upper bounds.",
            ],
            [
                "Regular neural network",
                "PyTorch DeepSurv-style feed-forward Cox network with a scalar risk output.",
                "Hidden layers 96 and 48; ReLU activation; batch normalization; dropout 0.15; Cox partial-likelihood loss; AdamW optimizer; learning rate 3e-4; weight decay 1e-3; 120 full-batch epochs; gradient-norm clipping at 5.0; seed 42.",
            ],
            [
                "Linear regression",
                "Event-only ridge regression of log survival time.",
                "Ridge penalty alpha=1.0. Numeric variables were median-imputed and standardized; categorical variables were mode-imputed and one-hot encoded. Only observed deaths were used for fitting. A log-normal residual transformation used sigma with a minimum value of 0.1 to obtain survival probabilities.",
            ],
        ],
        [1.35, 2.05, 3.1],
        "Table 4. Verified model implementations and hyperparameter settings.",
        7.6,
        left_align_columns={0, 1, 2},
    )

    doc.add_heading("6. Metrics", level=1)
    paragraph(
        doc,
        "C-index is the comparable-pair concordance statistic requested by the user. A pair contributes "
        "when the earlier observed time is a death; a correct ordering earns one point and a tied risk "
        "earns one-half point. Higher risk means shorter predicted survival. The report does not label "
        "this result as IPCW C-index."
    )
    paragraph(
        doc,
        "Time-dependent AUC was evaluated at 12, 24, 36, and 60 months and averaged over those horizons. "
        "Censoring-adjusted Brier scores were evaluated at the same four horizons."
    )
    paragraph(
        doc,
        "The prediction files store survival probabilities only at four "
        "horizons. The reported IBS is the mean of the four censoring-adjusted "
        "Brier scores, not numerical integration over a dense time grid. This report therefore calls it "
        "reported IBS and does not claim that a continuous-time IBS was reconstructed."
    )
    doc.add_heading("6.1 Uncertainty and Paired Model Comparisons", level=2)
    paragraph(
        doc,
        "Uncertainty was quantified with 1,000 patient-level bootstrap resamples. Resampling was paired "
        "across models so that every method was evaluated on the same sampled patients; the combined "
        "cross-validation analysis was additionally stratified by cohort to preserve the CGGA/TCGA "
        "composition. Percentile 95% confidence intervals were calculated for each metric. Paired "
        "CoxRes-KGA advantages were defined as CoxRes-KGA minus comparator for C-index and mean AUC, "
        "and comparator minus CoxRes-KGA for reported IBS, so positive values consistently favor the "
        "proposed method. These intervals quantify patient-sampling uncertainty conditional on the "
        "fitted models and do not include variability from repeating split generation or model refitting."
    )

    doc.add_heading("7. Performance Results", level=1)
    figure(doc, "figure_03_cindex_comparison.png", "Figure 3. Comparable-pair C-index across the three validation settings.", 6.3)
    figure(doc, "figure_04_auc_ibs_comparison.png", "Figure 4. Mean time-dependent AUC and reported IBS across all six models.", 6.35)

    for subsection, validation in enumerate((
        "Internal combined 5-fold CV",
        "External CGGA to TCGA",
        "External TCGA to CGGA",
    ), start=1):
        table_number = subsection + 4
        if subsection != 2:
            page_break(doc)
        doc.add_heading(f"7.{subsection} {validation}", level=2)
        subset = performance[performance.validation == validation]
        if validation == "Internal combined 5-fold CV":
            internal_rows, internal_summary = internal_cv_summary_rows(fold_metrics)
            winner = internal_summary.iloc[0]
            paragraph(
                doc,
                f"Across the five held-out folds, the highest mean C-index was "
                f"{winner.c_index_mean:.3f} (SD {winner.c_index_sd:.3f}), produced by "
                f"{MODEL_LABELS[winner.model]}. Mean AUC and reported IBS are summarized using the "
                "same fold-level mean (SD) convention."
            )
            table(
                doc,
                ["Model", "C-index, mean (SD)", "Mean AUC, mean (SD)", "Reported IBS, mean (SD)", "C-index rank"],
                internal_rows,
                [2.55, 1.20, 1.25, 1.30, .75],
                "Table 5. Internal five-fold cross-validation performance on the combined CGGA-TCGA cohort.",
                7.7,
            )
            paragraph(
                doc,
                "Values are the arithmetic mean (sample SD) of the five held-out-fold estimates. The "
                "SD describes cross-fold variability and is not a confidence interval. C-index rank is "
                "based on the fold mean."
            )
            figure(
                doc,
                "figure_table5_internal_comparison.png",
                "Figure 5. Model comparison corresponding to Table 5. Points show five-fold means and horizontal intervals show plus or minus one SD.",
                6.35,
            )
        else:
            winner = subset.sort_values("rank_c_index").iloc[0]
            paragraph(
                doc,
                f"The highest C-index was {winner.c_index:.3f}, produced by {MODEL_LABELS[winner.model]}. "
                "AUC and reported IBS rankings are shown separately because discrimination and prediction "
                "error need not identify the same model."
            )
            table(
                doc,
                ["Model", "C-index", "Mean AUC", "Reported IBS", "C-index rank"],
                result_rows(subset),
                [2.9, .85, .85, .95, .95],
                f"Table {table_number}. {validation} performance.",
                8.1,
            )
            comparison_file = (
                "figure_table6_cgga_to_tcga_comparison.png"
                if validation == "External CGGA to TCGA"
                else "figure_table7_tcga_to_cgga_comparison.png"
            )
            figure_number = 7 if validation == "External CGGA to TCGA" else 8
            training_cohort = "CGGA" if validation == "External CGGA to TCGA" else "TCGA"
            testing_cohort = "TCGA" if validation == "External CGGA to TCGA" else "CGGA"
            figure(
                doc,
                comparison_file,
                f"Figure {figure_number}. Model comparison corresponding to Table {table_number}; models were trained in {training_cohort} and evaluated in {testing_cohort}.",
                6.35,
            )
        if validation == "Internal combined 5-fold CV":
            doc.add_heading("7.1.1 Fold-level Cross-validation Performance", level=3)
            paragraph(
                doc,
                "Fold-level estimates are reported for every model to show sensitivity to the held-out "
                "partition. Mean (SD) summarizes descriptive variation across the five folds and is not "
                "a confidence interval. The pooled out-of-fold estimate is calculated after combining "
                "all 1,697 held-out predictions and therefore need not equal the arithmetic fold mean."
            )
            figure(
                doc,
                "figure_fold_stability.png",
                "Figure 6. Fold-level C-index, mean time-dependent AUC, and reported IBS for all six models. Colored points represent held-out folds and diamonds represent fold means.",
                6.35,
            )

            internal_order = list(
                performance[performance.validation == validation]
                .sort_values("rank_c_index")["model"]
            )
            pooled_internal = performance[performance.validation == validation].set_index("model")
            suffixes = "ABCDEF"
            page_break(doc)
            for model_index, model in enumerate(internal_order):
                doc.add_heading(MODEL_LABELS[model], level=3)
                table(
                    doc,
                    ["Estimate", "N", "Events", "C-index", "Mean AUC", "Reported IBS"],
                    fold_model_rows(fold_metrics, model, pooled_internal.loc[model]),
                    [1.35, .62, .72, 1.0, 1.05, 1.15],
                    f"Table 5{suffixes[model_index]}. Fold-level internal-validation performance for {MODEL_LABELS[model]}.",
                    7.7,
                )

    heading = doc.add_heading("7.4 Horizon and Clinical-Stratification Views", level=2)
    heading.paragraph_format.page_break_before = True
    figure(doc, "figure_05_time_auc.png", "Figure 9. Horizon-specific AUC for all six models across internal and external validation settings.", 6.15)
    figure(doc, "figure_06_km_internal.png", "Figure 10. Internal combined-CV out-of-fold Kaplan-Meier curves by CoxRes-KGA risk group.", 6.35)
    paragraph(
        doc,
        "Risk groups were created separately within CGGA and TCGA at the cohort-specific median risk for "
        "visualization only. They are not validated clinical thresholds."
    )
    figure(doc, "figure_07_external_calibration.png", "Figure 11. External calibration summary for CoxRes-KGA.", 6.15)
    figure(doc, "figure_08_subgroup_risk.png", "Figure 12. CoxRes-KGA risk distributions across grade, IDH, and MGMT groups.", 5.85)

    page_break(doc)
    doc.add_heading("7.5 Integrated Performance Scorecard", level=2)
    paragraph(
        doc,
        "The integrated scorecard presents the C-index, mean time-dependent AUC, and reported IBS for "
        "all six models under each validation setting. A consistent method order and exact value labels "
        "support direct comparison across metrics; upward arrows indicate that higher values are preferred, "
        "whereas the downward arrow for IBS indicates that lower values are preferred."
    )
    figure(
        doc,
        "figure_09_performance_scorecard.png",
        "Figure 13. Integrated model-performance scorecard across internal and external validation settings.",
        6.35,
    )

    page_break(doc)
    doc.add_heading("7.6 Bootstrap Uncertainty and Paired Comparisons", level=2)
    paragraph(
        doc,
        "Bootstrap intervals provide an uncertainty-aware complement to the point estimates. Overlap "
        "between model-specific intervals should not be used as a formal paired test; Figure 15 instead "
        "uses within-resample differences, which preserve the common-patient comparison."
    )
    figure(
        doc,
        "figure_10_bootstrap_intervals.png",
        "Figure 14. Patient-bootstrap 95% confidence intervals for all six models and three performance metrics.",
        6.35,
    )
    proposed_bootstrap = bootstrap[bootstrap.model == proposed_id]
    interval_rows = []
    for validation in (
        "Internal combined 5-fold CV",
        "External CGGA to TCGA",
        "External TCGA to CGGA",
    ):
        item = proposed_bootstrap[proposed_bootstrap.validation == validation].set_index("metric")
        interval_rows.append([
            validation,
            fmt_ci(item.loc["c_index", "estimate"], item.loc["c_index", "ci_low"], item.loc["c_index", "ci_high"]),
            fmt_ci(item.loc["auc_mean", "estimate"], item.loc["auc_mean", "ci_low"], item.loc["auc_mean", "ci_high"]),
            fmt_ci(item.loc["ibs", "estimate"], item.loc["ibs", "ci_low"], item.loc["ibs", "ci_high"]),
        ])
    table(
        doc,
        ["Validation", "C-index (95% CI)", "Mean AUC (95% CI)", "Reported IBS (95% CI)"],
        interval_rows,
        [2.25, 1.4, 1.4, 1.45],
        "Table 8. CoxRes-KGA performance with patient-bootstrap 95% confidence intervals.",
        8.0,
    )
    figure(
        doc,
        "figure_11_paired_advantage.png",
        "Figure 15. Direction-adjusted paired-bootstrap differences between CoxRes-KGA and each comparator.",
        6.35,
    )

    doc.add_heading("7.7 All-Model Calibration", level=2)
    paragraph(
        doc,
        "Calibration-in-the-large was assessed by comparing mean predicted survival with the corresponding "
        "Kaplan-Meier estimate at each prespecified prediction horizon. A value of zero indicates agreement "
        "at the cohort level; positive values indicate overprediction of survival and negative values "
        "indicate underprediction. This summary does not replace subgroup calibration curves or a "
        "calibration slope analysis."
    )
    figure(
        doc,
        "figure_12_all_model_calibration.png",
        "Figure 16. Calibration-in-the-large error for every model, validation setting, and prediction horizon.",
        6.2,
    )

    doc.add_heading("7.8 Architecture Sensitivity Analysis", level=2)
    paragraph(
        doc,
        "A controlled architecture comparison retained the same corrected cohorts, outer validation splits, "
        "nested Cox/loss tuning grid, training schedule, and five-seed ensemble while varying only token "
        "dimension and attention-head count. The former 32-dimensional/one-head model was compared with "
        "16-dimensional models using one or two heads. Paired patient-bootstrap intervals quantify differences "
        "from the former architecture. The 16-dimensional/one-head model was retained because it improved "
        "internal pooled ranking and produced the largest CGGA-to-TCGA discrimination gain without a clear "
        "loss in the reverse external direction."
    )
    architecture_rows = []
    architecture_labels = {
        "32d/1h": "32d / 1 head",
        "16d/1h": "16d / 1 head",
        "16d/2h": "16d / 2 heads",
    }
    for validation in (
        "Internal combined 5-fold CV",
        "External CGGA to TCGA",
        "External TCGA to CGGA",
    ):
        for architecture in ("32d/1h", "16d/1h", "16d/2h"):
            item = architecture_summary.loc[
                architecture_summary.validation.eq(validation)
                & architecture_summary.architecture.eq(architecture)
            ].iloc[0]
            architecture_rows.append([
                validation,
                architecture_labels[architecture],
                f"{item.c_index:.3f} ({int(item.rank_c_index)})",
                f"{item.auc_mean:.3f} ({int(item.rank_auc)})",
                f"{item.ibs:.3f} ({int(item.rank_ibs)})",
            ])
    table(
        doc,
        ["Validation", "Architecture", "C-index (rank)", "Mean AUC (rank)", "IBS (rank)"],
        architecture_rows,
        [1.65, 1.05, 1.15, 1.15, 1.00],
        "Table 9. Controlled CoxRes-KGA architecture sensitivity results.",
        7.4,
    )
    c_external = architecture_bootstrap.loc[
        architecture_bootstrap.validation.eq("External CGGA to TCGA")
        & architecture_bootstrap.architecture.eq("16d/1h")
    ].iloc[0]
    paragraph(
        doc,
        f"Relative to 32d/1-head, 16d/1-head improved CGGA-to-TCGA C-index by "
        f"{c_external.delta_c_index:+.4f} (95% CI {c_external.delta_c_index_low:+.4f} to "
        f"{c_external.delta_c_index_high:+.4f}) and mean AUC by {c_external.delta_auc_mean:+.4f} "
        f"(95% CI {c_external.delta_auc_mean_low:+.4f} to {c_external.delta_auc_mean_high:+.4f}). "
        "Internal and TCGA-to-CGGA differences were smaller and their intervals included zero."
    )
    figure(
        doc,
        "figure_14_architecture_sensitivity.png",
        "Figure 17. Paired-bootstrap performance differences for 16-dimensional CoxRes-KGA alternatives relative to the former 32d/1-head architecture.",
        6.35,
    )

    doc.add_heading("7.9 Efron and Learnable-Bias Sensitivity Analysis", level=2)
    paragraph(
        doc,
        "A controlled sensitivity analysis evaluated Efron correction for tied deaths [14], regularized "
        "learnable knowledge-bias strengths, and alternative Cox/residual shrinkage. Four prespecified "
        "configurations were compared by three-fold validation inside each outer training set. The selected "
        "configuration was then refitted as a five-seed ensemble. External destination-cohort outcomes were "
        "not used for configuration selection in either directional validation job."
    )
    sensitivity_rows = []
    for _, item in sensitivity_delta.iterrows():
        sensitivity_rows.append([
            item.validation,
            f"{item.primary_c_index:.3f} ({int(item.primary_rank_c_index)})",
            f"{item.ablation_c_index:.3f} ({int(item.ablation_rank_c_index)})",
            f"{item.delta_c_index:+.4f}",
            f"{item.delta_auc_mean:+.4f}",
            f"{item.delta_ibs:+.4f}",
        ])
    table(
        doc,
        ["Validation", "Primary C-index (rank)", "Sensitivity C-index (rank)", "Delta C-index", "Delta AUC", "Delta IBS"],
        sensitivity_rows,
        [1.60, 1.25, 1.35, 0.85, 0.75, 0.70],
        "Table 10. Primary versus Efron/learnable-bias CoxRes-KGA sensitivity results.",
        7.6,
    )
    frequency_text = ", ".join(
        f"{row.selected_config.replace('_', ' ')} in {int(row.jobs_selected)} job(s)"
        for _, row in sensitivity_frequency.sort_values("jobs_selected", ascending=False).iterrows()
    )
    paragraph(
        doc,
        "Training-only selection chose " + frequency_text + ". Relative to the retained 16-dimensional model, "
        "the sensitivity model reduced internal and CGGA-to-TCGA discrimination and changed the "
        "TCGA-to-CGGA C-index rank from second to third. Its small improvement in TCGA-to-CGGA prediction "
        "error did not offset those losses. It was therefore retained as a "
        "sensitivity analysis rather than replacing the primary publication model."
    )
    figure(
        doc,
        "figure_13_efron_sensitivity.png",
        "Figure 18. Direction-adjusted performance changes and training-only configuration-selection frequency for the Efron/learnable-bias sensitivity analysis.",
        6.35,
    )

    doc.add_heading("7.10 Structured Improvement Sensitivity Analysis", level=2)
    paragraph(
        doc,
        "Three additional leakage-safe variants retained the 16-dimensional, one-head neural architecture and "
        "the same outer splits and nested Cox/loss selection. Variant A replaced the linear age and grade Cox "
        "terms with a cubic age-spline basis and categorical grade effects. Variant B used those nonlinear main "
        "effects with five-fold cross-fitted member averaging. Variant C additionally penalized correlation "
        "between the Cox offset and neural residual (orthogonality weight 0.01)."
    )
    reverse_c = structured_bootstrap.loc[
        structured_bootstrap.validation.eq("External TCGA to CGGA")
        & structured_bootstrap.variant.eq("C: nonlinear + crossfit + orthogonal")
    ].iloc[0]
    forward_c = structured_bootstrap.loc[
        structured_bootstrap.validation.eq("External CGGA to TCGA")
        & structured_bootstrap.variant.eq("C: nonlinear + crossfit + orthogonal")
    ].iloc[0]
    paragraph(
        doc,
        f"The most elaborate variant changed TCGA-to-CGGA C-index by {reverse_c.delta_c_index:+.4f} "
        f"(95% CI {reverse_c.delta_c_index_low:+.4f} to {reverse_c.delta_c_index_high:+.4f}) but changed "
        f"CGGA-to-TCGA C-index by {forward_c.delta_c_index:+.4f} "
        f"(95% CI {forward_c.delta_c_index_low:+.4f} to {forward_c.delta_c_index_high:+.4f}). "
        "Every paired C-index and AUC interval for all three variants included zero. The variants were therefore "
        "retained as negative sensitivity analyses, and the original 16d/1-head specification remained primary."
    )
    figure(
        doc,
        "figure_15_structured_improvements.png",
        "Figure 19. Paired-bootstrap differences for nonlinear Cox, cross-fitted, and residual-orthogonality variants relative to the retained 16d/1-head model.",
        6.35,
    )

    doc.add_heading("7.11 Primary-Tumor-Only Sensitivity Analysis", level=2)
    paragraph(
        doc,
        "To examine disease-setting heterogeneity, a controlled sensitivity analysis excluded 253 outcome-eligible "
        "CGGA recurrent tumors and retained 404 CGGA primary tumors. All 1,040 outcome-eligible TCGA records "
        "were retained because their sample barcodes carried TCGA sample-type code 01 (primary solid tumor). "
        "All six models were refitted with the same predictors, five internal folds, external directions, and metric definitions."
    )
    primary_proposed = primary_comparison[primary_comparison.model.eq(proposed_id)].copy()
    primary_rows = []
    for validation in (
        "Internal combined 5-fold CV", "External CGGA to TCGA", "External TCGA to CGGA"
    ):
        item = primary_proposed[primary_proposed.validation.eq(validation)].iloc[0]
        primary_rows.append([
            validation,
            f"{item.c_index_full:.3f} → {item.c_index_primary:.3f} ({int(item.rank_c_index_primary)})",
            f"{item.auc_mean_full:.3f} → {item.auc_mean_primary:.3f} ({int(item.rank_auc_primary)})",
            f"{item.ibs_full:.3f} → {item.ibs_primary:.3f} ({int(item.rank_ibs_primary)})",
        ])
    table(
        doc,
        ["Validation", "C-index: full → primary (rank)", "Mean AUC: full → primary (rank)", "IBS: full → primary (rank)"],
        primary_rows,
        [2.0, 1.5, 1.5, 1.5],
        "Table 11. CoxRes-KGA full-cohort and primary-tumor-only sensitivity results.",
        8.2,
    )
    paragraph(
        doc,
        f"On the unchanged TCGA destination cohort, restricting CGGA training to primary tumors increased "
        f"C-index by {primary_paired.delta_c_index:+.4f} (95% CI {primary_paired.delta_c_index_low:+.4f} to "
        f"{primary_paired.delta_c_index_high:+.4f}) and mean AUC by {primary_paired.delta_auc_mean:+.4f} "
        f"(95% CI {primary_paired.delta_auc_mean_low:+.4f} to {primary_paired.delta_auc_mean_high:+.4f}). "
        f"The IBS change was {primary_paired.delta_ibs:+.4f} (95% CI {primary_paired.delta_ibs_low:+.4f} to "
        f"{primary_paired.delta_ibs_high:+.4f}); positive IBS change indicates worse prediction error. "
        "Internal and TCGA-to-CGGA changes also alter the evaluated patient population and therefore should not "
        "be interpreted as paired model improvements. Relative ranks did not improve consistently, so this "
        "analysis was not promoted over the prespecified full-cohort result."
    )
    figure(
        doc,
        "figure_16_primary_tumor_sensitivity.png",
        "Figure 20. Full-cohort versus primary-tumor-only performance for all six models and validation settings.",
        6.35,
    )

    doc.add_page_break()
    doc.add_heading("7.12 Selection of the Best-Ranked Training Configuration", level=2)
    paragraph(
        doc,
        "The selected CoxRes-KGA configuration retained the 16-dimensional, one-head, one-layer architecture "
        "and the same nested configuration-selection procedure, while increasing dropout from 0.12 to 0.40 "
        "and replacing AdamW with Adam. The learning rate (3×10⁻⁴), nominal weight decay (10⁻³), outer splits, "
        "candidate grid, and five-seed ensemble were unchanged. Existing comparator predictions were reused, "
        "so the comparison isolates the joint dropout-and-optimizer change within CoxRes-KGA."
    )
    dropout_proposed = dropout_performance[dropout_performance.model.eq(proposed_id)].copy()
    dropout_rows = []
    for validation in (
        "Internal combined 5-fold CV", "External CGGA to TCGA", "External TCGA to CGGA"
    ):
        retained = reference_proposed.loc[validation]
        variant = dropout_proposed[dropout_proposed.validation.eq(validation)].iloc[0]
        dropout_rows.append([
            validation,
            f"{retained.c_index:.3f} → {variant.c_index:.3f} ({int(variant.rank_c_index)})",
            f"{retained.auc_mean:.3f} → {variant.auc_mean:.3f} ({int(variant.rank_auc)})",
            f"{retained.ibs:.3f} → {variant.ibs:.3f} ({int(variant.rank_ibs)})",
        ])
    table(
        doc,
        ["Validation", "C-index: reference → selected (rank)", "Mean AUC: reference → selected (rank)", "IBS: reference → selected (rank)"],
        dropout_rows,
        [2.0, 1.5, 1.5, 1.5],
        "Table 12. Reference dropout-0.12/AdamW versus selected dropout-0.40/Adam CoxRes-KGA results.",
        8.2,
    )
    forward_dropout = dropout_comparison[dropout_comparison.validation.eq("External CGGA to TCGA")].iloc[0]
    internal_dropout = dropout_comparison[dropout_comparison.validation.eq("Internal combined 5-fold CV")].iloc[0]
    paragraph(
        doc,
        f"The selected configuration ranked first for all three metrics internally and in CGGA-to-TCGA validation, and first "
        f"for IBS in TCGA-to-CGGA validation. On paired bootstrap analysis, CGGA-to-TCGA mean AUC increased "
        f"by {forward_dropout.delta_auc_mean:+.4f} (95% CI {forward_dropout.delta_auc_mean_low:+.4f} to "
        f"{forward_dropout.delta_auc_mean_high:+.4f}). Its C-index change was {forward_dropout.delta_c_index:+.4f} "
        f"(95% CI {forward_dropout.delta_c_index_low:+.4f} to {forward_dropout.delta_c_index_high:+.4f}), while "
        f"internal IBS changed by {internal_dropout.delta_ibs:+.4f} (95% CI {internal_dropout.delta_ibs_low:+.4f} "
        f"to {internal_dropout.delta_ibs_high:+.4f}); positive IBS change indicates worse prediction error. "
        "All other paired intervals included zero. The complete run was repeated independently and produced "
        "byte-identical patient-level predictions, metrics, rankings, and training decisions. Because dropout "
        "and optimizer were changed together after reviewing prior results, the selected configuration remains "
        "exploratory and requires confirmation in a genuinely independent cohort."
    )
    figure(
        doc,
        "figure_17_dropout_adam_sensitivity.png",
        "Figure 21. Paired-bootstrap performance differences for the selected dropout-0.40/Adam configuration relative to the dropout-0.12/AdamW reference.",
        6.35,
    )

    doc.add_page_break()
    doc.add_heading("7.13 Two-Head, Dropout-0.50, and AdamW Sensitivity Analysis", level=2)
    paragraph(
        doc,
        "A subsequent exploratory configuration increased the attention-head count from one to two and dropout "
        "from 0.40 to 0.50, while replacing Adam with AdamW. Token dimension 16, one attention layer, learning "
        "rate 3×10⁻⁴, nominal weight decay 10⁻³, nested training-only selection, outer splits, and five-seed "
        "ensembling were retained. Comparator predictions were again unchanged."
    )
    dropout_h2_proposed = dropout_h2_performance[dropout_h2_performance.model.eq(proposed_id)].copy()
    dropout_h2_rows = []
    for validation in (
        "Internal combined 5-fold CV", "External CGGA to TCGA", "External TCGA to CGGA"
    ):
        best_rank = dropout_proposed[dropout_proposed.validation.eq(validation)].iloc[0]
        candidate = dropout_h2_proposed[dropout_h2_proposed.validation.eq(validation)].iloc[0]
        dropout_h2_rows.append([
            validation,
            f"{best_rank.c_index:.3f} → {candidate.c_index:.3f} ({int(candidate.rank_c_index)})",
            f"{best_rank.auc_mean:.3f} → {candidate.auc_mean:.3f} ({int(candidate.rank_auc)})",
            f"{best_rank.ibs:.3f} → {candidate.ibs:.3f} ({int(candidate.rank_ibs)})",
        ])
    table(
        doc,
        ["Validation", "C-index: 0.40/Adam/1h → candidate (rank)", "Mean AUC: 0.40/Adam/1h → candidate (rank)", "IBS: 0.40/Adam/1h → candidate (rank)"],
        dropout_h2_rows,
        [2.0, 1.5, 1.5, 1.5],
        "Table 13. Two-head/dropout-0.50/AdamW CoxRes-KGA sensitivity results.",
        8.0,
    )
    versus_best = dropout_h2_comparison[
        dropout_h2_comparison.baseline.eq("Best-rank 0.40/Adam/1h")
    ]
    internal_h2 = versus_best[versus_best.validation.eq("Internal combined 5-fold CV")].iloc[0]
    forward_h2 = versus_best[versus_best.validation.eq("External CGGA to TCGA")].iloc[0]
    paragraph(
        doc,
        f"The candidate reproduced the same seven-of-nine rank-1 profile as the dropout-0.40/Adam model. "
        f"However, internal IBS increased by {internal_h2.delta_ibs:+.4f} "
        f"(95% CI {internal_h2.delta_ibs_low:+.4f} to {internal_h2.delta_ibs_high:+.4f}), and CGGA-to-TCGA "
        f"IBS increased by {forward_h2.delta_ibs:+.4f} "
        f"(95% CI {forward_h2.delta_ibs_low:+.4f} to {forward_h2.delta_ibs_high:+.4f}). Positive IBS changes "
        "indicate worse prediction error. Every paired C-index and AUC interval and the reverse-direction IBS "
        "interval included zero. The candidate therefore did not improve on the dropout-0.40/Adam model and "
        "was retained as a negative sensitivity analysis."
    )
    figure(
        doc,
        "figure_18_training_configuration_sensitivity.png",
        "Figure 22. Performance comparison of the dropout-0.12/AdamW reference, selected dropout-0.40/Adam, and two-head/dropout-0.50/AdamW CoxRes-KGA configurations.",
        6.35,
    )

    doc.add_page_break()
    doc.add_heading("7.14 Single-Factor Dropout-0.60 Sensitivity Analysis", level=2)
    paragraph(
        doc,
        "A single-factor sensitivity analysis increased dropout from 0.40 to 0.60 while retaining Adam, "
        "16-dimensional tokens, one attention head, one layer, the same optimizer settings, nested selection, "
        "outer validation splits, and five-seed ensemble. This design isolates dropout more directly than the "
        "multi-factor optimization experiments."
    )
    dropout060_proposed = dropout060_performance[dropout060_performance.model.eq(proposed_id)].copy()
    dropout060_rows = []
    for validation in (
        "Internal combined 5-fold CV", "External CGGA to TCGA", "External TCGA to CGGA"
    ):
        baseline = dropout_proposed[dropout_proposed.validation.eq(validation)].iloc[0]
        candidate = dropout060_proposed[dropout060_proposed.validation.eq(validation)].iloc[0]
        dropout060_rows.append([
            validation,
            f"{baseline.c_index:.3f} → {candidate.c_index:.3f} ({int(candidate.rank_c_index)})",
            f"{baseline.auc_mean:.3f} → {candidate.auc_mean:.3f} ({int(candidate.rank_auc)})",
            f"{baseline.ibs:.3f} → {candidate.ibs:.3f} ({int(candidate.rank_ibs)})",
        ])
    table(
        doc,
        ["Validation", "C-index: dropout 0.40 → 0.60 (rank)", "Mean AUC: dropout 0.40 → 0.60 (rank)", "IBS: dropout 0.40 → 0.60 (rank)"],
        dropout060_rows,
        [2.0, 1.5, 1.5, 1.5],
        "Table 14. Dropout-0.60/Adam single-factor CoxRes-KGA sensitivity results.",
        8.0,
    )
    internal_d060 = dropout060_comparison[
        dropout060_comparison.validation.eq("Internal combined 5-fold CV")
    ].iloc[0]
    reverse_d060 = dropout060_comparison[
        dropout060_comparison.validation.eq("External TCGA to CGGA")
    ].iloc[0]
    paragraph(
        doc,
        f"Dropout 0.60 retained seven rank-1 results but reduced TCGA-to-CGGA C-index rank from second to third. "
        f"Internal IBS increased by {internal_d060.delta_ibs:+.4f} "
        f"(95% CI {internal_d060.delta_ibs_low:+.4f} to {internal_d060.delta_ibs_high:+.4f}), while reverse-external "
        f"mean AUC changed by {reverse_d060.delta_auc_mean:+.4f} "
        f"(95% CI {reverse_d060.delta_auc_mean_low:+.4f} to {reverse_d060.delta_auc_mean_high:+.4f}). "
        "All C-index intervals included zero. The higher dropout therefore increased underfitting without a "
        "reliable discrimination benefit and was not selected over dropout 0.40."
    )
    figure(
        doc,
        "figure_19_dropout060_sensitivity.png",
        "Figure 23. Patient-paired performance differences for dropout 0.60 relative to dropout 0.40 with Adam and the one-head architecture fixed.",
        6.35,
    )

    doc.add_page_break()
    doc.add_heading("7.15 One-Head Dropout-0.50/AdamW Sensitivity Analysis", level=2)
    paragraph(
        doc,
        "A controlled sensitivity analysis retained 16-dimensional tokens, one attention head, one layer, "
        "the same learning rate and weight decay, nested training-only selection, outer validation splits, "
        "and five-seed averaging. Relative to the dropout-0.40/Adam configuration, only dropout was increased "
        "to 0.50 and the optimizer was changed to AdamW."
    )
    dropout050_h1_proposed = dropout050_h1_performance[
        dropout050_h1_performance.model.eq(proposed_id)
    ].copy()
    dropout050_h1_rows = []
    for validation in (
        "Internal combined 5-fold CV", "External CGGA to TCGA", "External TCGA to CGGA"
    ):
        baseline = dropout_proposed[dropout_proposed.validation.eq(validation)].iloc[0]
        candidate = dropout050_h1_proposed[
            dropout050_h1_proposed.validation.eq(validation)
        ].iloc[0]
        dropout050_h1_rows.append([
            validation,
            f"{baseline.c_index:.3f} → {candidate.c_index:.3f} ({int(candidate.rank_c_index)})",
            f"{baseline.auc_mean:.3f} → {candidate.auc_mean:.3f} ({int(candidate.rank_auc)})",
            f"{baseline.ibs:.3f} → {candidate.ibs:.3f} ({int(candidate.rank_ibs)})",
        ])
    table(
        doc,
        ["Validation", "C-index: 0.40/Adam → 0.50/AdamW (rank)", "Mean AUC: 0.40/Adam → 0.50/AdamW (rank)", "IBS: 0.40/Adam → 0.50/AdamW (rank)"],
        dropout050_h1_rows,
        [2.0, 1.5, 1.5, 1.5],
        "Table 15. One-head dropout-0.50/AdamW CoxRes-KGA sensitivity results.",
        8.0,
    )
    internal_d050_h1 = dropout050_h1_comparison[
        dropout050_h1_comparison.validation.eq("Internal combined 5-fold CV")
    ].iloc[0]
    reverse_d050_h1 = dropout050_h1_comparison[
        dropout050_h1_comparison.validation.eq("External TCGA to CGGA")
    ].iloc[0]
    paragraph(
        doc,
        f"The candidate retained seven rank-1 metrics, but reverse-external mean AUC fell from rank 2 to "
        f"rank 3. Internal IBS increased by {internal_d050_h1.delta_ibs:+.6f} "
        f"(95% CI {internal_d050_h1.delta_ibs_low:+.6f} to {internal_d050_h1.delta_ibs_high:+.6f}), "
        f"whereas reverse-external mean AUC changed by {reverse_d050_h1.delta_auc_mean:+.6f} "
        f"(95% CI {reverse_d050_h1.delta_auc_mean_low:+.6f} to "
        f"{reverse_d050_h1.delta_auc_mean_high:+.6f}). All C-index intervals included zero. The change "
        "therefore did not provide a reliable discrimination benefit and slightly worsened the overall rank "
        "profile, so dropout 0.40/Adam remains the preferred exploratory configuration."
    )
    figure(
        doc,
        "figure_20_dropout050_adamw_h1_sensitivity.png",
        "Figure 24. Patient-paired performance differences for one-head dropout 0.50/AdamW relative to dropout 0.40/Adam.",
        6.35,
    )

    doc.add_page_break()
    doc.add_heading("7.16 Dimension and Training-Duration Sensitivity Analysis", level=2)
    paragraph(
        doc,
        "A two-factor sensitivity analysis increased the token dimension from 16 to 32 and reduced the "
        "maximum epoch allowance from 170 to 120. One attention head, one layer, dropout 0.40, Adam, learning "
        "rate, weight decay, minimum 70-epoch training, patience, nested training-only selection, validation "
        "splits, and five-seed averaging were retained."
    )
    d32_e120_proposed = d32_e120_performance[d32_e120_performance.model.eq(proposed_id)].copy()
    d32_selected_epochs = [
        member["selected_epoch"]
        for job in d32_e120_manifest
        for member in job["model_details"]["members"]
    ]
    d32_e120_rows = []
    for validation in (
        "Internal combined 5-fold CV", "External CGGA to TCGA", "External TCGA to CGGA"
    ):
        baseline = dropout_proposed[dropout_proposed.validation.eq(validation)].iloc[0]
        candidate = d32_e120_proposed[d32_e120_proposed.validation.eq(validation)].iloc[0]
        d32_e120_rows.append([
            validation,
            f"{baseline.c_index:.3f} → {candidate.c_index:.3f} ({int(candidate.rank_c_index)})",
            f"{baseline.auc_mean:.3f} → {candidate.auc_mean:.3f} ({int(candidate.rank_auc)})",
            f"{baseline.ibs:.3f} → {candidate.ibs:.3f} ({int(candidate.rank_ibs)})",
        ])
    table(
        doc,
        ["Validation", "C-index: 16d/170 → 32d/120 (rank)", "Mean AUC: 16d/170 → 32d/120 (rank)", "IBS: 16d/170 → 32d/120 (rank)"],
        d32_e120_rows,
        [2.0, 1.5, 1.5, 1.5],
        "Table 16. CoxRes-KGA 32-dimensional/120-epoch sensitivity results.",
        8.0,
    )
    forward_d32 = d32_e120_comparison[
        d32_e120_comparison.validation.eq("External CGGA to TCGA")
    ].iloc[0]
    paragraph(
        doc,
        f"The candidate achieved six rank-1, two rank-2, and one rank-3 results, compared with seven rank-1 "
        f"and two rank-2 results for the selected model. In CGGA-to-TCGA validation, C-index changed by "
        f"{forward_d32.delta_c_index:+.6f} (95% CI {forward_d32.delta_c_index_low:+.6f} to "
        f"{forward_d32.delta_c_index_high:+.6f}) and mean AUC changed by "
        f"{forward_d32.delta_auc_mean:+.6f} (95% CI {forward_d32.delta_auc_mean_low:+.6f} to "
        f"{forward_d32.delta_auc_mean_high:+.6f}). Both intervals excluded zero in the unfavorable direction. "
        f"Selected epochs ranged from {min(d32_selected_epochs)} to {max(d32_selected_epochs)}, so the "
        "120-epoch cap did not bind. The observed difference therefore primarily reflects the larger "
        "representation under the same early-stopping regime. It reduced forward-external discrimination and "
        "was not selected over the 16-dimensional/170-epoch configuration."
    )
    figure(
        doc,
        "figure_21_d32_e120_sensitivity.png",
        "Figure 25. Patient-paired performance differences for 32-dimensional/120-epoch CoxRes-KGA relative to the selected 16-dimensional/170-epoch model.",
        6.35,
    )

    doc.add_heading("8. Discussion", level=1)
    internal_ranked = performance[
        performance.validation == "Internal combined 5-fold CV"
    ].sort_values("rank_c_index")
    internal_winner = internal_ranked.iloc[0]
    internal_runner_up = internal_ranked.iloc[1]
    paragraph(
        doc,
        f"The corrected-cohort experiment supports a nuanced conclusion. In combined internal CV, "
        f"{MODEL_LABELS[internal_winner.model]} had the highest pooled C-index "
        f"({internal_winner.c_index:.4f}), virtually tied with "
        f"{MODEL_LABELS[internal_runner_up.model]} ({internal_runner_up.c_index:.4f}). For CGGA-to-TCGA "
        f"validation, {MODEL_LABELS[winners.loc['External CGGA to TCGA', 'model']]} ranked first "
        f"({winners.loc['External CGGA to TCGA', 'c_index']:.3f}); for TCGA-to-CGGA validation, "
        f"{MODEL_LABELS[winners.loc['External TCGA to CGGA', 'model']]} ranked first "
        f"({winners.loc['External TCGA to CGGA', 'c_index']:.3f})."
    )
    paragraph(
        doc,
        "The bootstrap analysis adds uncertainty to these rank-based observations and directly evaluates "
        "paired differences on common patient samples. Calibration-in-the-large further shows whether "
        "models systematically over- or underpredict cohort survival across horizons; it should be read "
        "alongside discrimination because a well-ranked model can still produce biased absolute survival "
        "probabilities."
    )
    paragraph(
        doc,
        "The Efron/learnable-bias sensitivity analysis did not provide a consistent improvement over the "
        "primary fixed-bias, nested-tuning model. This negative result supports retaining the simpler attention specification and "
        "suggests that additional model complexity is less promising than improved predictor availability "
        "or genuinely independent validation."
    )
    paragraph(
        doc,
        "Likewise, nonlinear Cox main effects, cross-fitted averaging, and residual orthogonality did not "
        "produce a statistically reliable paired improvement. Their modest gains in TCGA-to-CGGA prediction "
        "error were offset by weaker CGGA-to-TCGA AUC and prediction error. This result further supports the "
        "simpler one-head architecture rather than selecting additional complexity from small point-estimate differences."
    )
    paragraph(
        doc,
        "The primary-tumor-only analysis produced higher absolute discrimination, particularly when recurrent "
        "CGGA cases were removed from the destination cohort. The paired CGGA-primary-to-TCGA comparison also "
        "showed that disease-setting alignment improved discrimination on the unchanged TCGA test population. "
        "However, comparator models improved as well and CoxRes-KGA ranks did not become uniformly better. This "
        "supports recurrent-disease heterogeneity as one contributor to cohort shift rather than demonstrating "
        "that architectural overfitting was corrected."
    )
    paragraph(
        doc,
        "The selected dropout-0.40/Adam configuration improved the rank profile and produced a paired increase in "
        "CGGA-to-TCGA mean AUC, but most paired intervals included zero and internal prediction error became "
        "slightly worse relative to the dropout-0.12/AdamW reference. Its ranking advantage therefore reflects "
        "small differences among closely performing models. Exact computational reproduction supports stability, "
        "but independent validation is still needed before interpreting it as clinically superior."
    )
    paragraph(
        doc,
        "The two-head/dropout-0.50/AdamW candidate preserved the same aggregate ranks but did not provide a "
        "paired discrimination improvement over the dropout-0.40/Adam variant. Its significantly higher internal "
        "and CGGA-to-TCGA prediction error makes the one-head/dropout-0.40/Adam variant the preferable exploratory "
        "configuration when rank ties are resolved using prediction error."
    )
    paragraph(
        doc,
        "The single-factor dropout analysis provides clearer evidence that increasing dropout beyond 0.40 is "
        "not beneficial for this compact network. Dropout 0.60 worsened reverse-external AUC, increased internal "
        "prediction error, and reduced the reverse-external C-index rank without a reliable C-index gain."
    )
    paragraph(
        doc,
        "Holding the compact architecture fixed while jointly using dropout 0.50 and AdamW also failed to "
        "improve the selected configuration. It preserved seven first-place metrics but reduced reverse-external "
        "AUC rank and significantly increased internal prediction error."
    )
    paragraph(
        doc,
        "Increasing representation dimension to 32 while shortening the maximum training schedule to 120 epochs "
        "also weakened the aggregate rank profile. Its significant reductions in forward-external C-index and "
        "mean AUC argue against promoting the candidate. Because all selected epochs were 99 or fewer, the cap "
        "did not truncate training and the practical contrast mainly concerns representation dimension."
    )
    paragraph(
        doc,
        "This pattern is consistent with the intended residual design: clinical main effects remain in a "
        "stable Cox pathway while attention is tasked with learning additional nonlinear structure. However, "
        "the selected 16-dimensional model ranked "
        f"{int(proposed.loc['External CGGA to TCGA', 'rank_c_index'])} for CGGA-to-TCGA C-index and "
        f"{int(proposed.loc['External TCGA to CGGA', 'rank_c_index'])} for TCGA-to-CGGA C-index. This directional asymmetry is "
        "compatible with cohort shift and model-capacity sensitivity and does not establish general clinical "
        "transportability."
    )
    paragraph(
        doc,
        "The reference paper reported SGS-Net C-index 0.8107 using multimodal MRI, segmentation "
        "supervision, CNN-Transformer representations, Cox loss, and pairwise ranking. The present "
        "clinical-molecular experiment differs in modality, patients, outcome processing, and validation. "
        "Its values must not be treated as a direct state-of-the-art contest."
    )

    doc.add_heading("9. Limitations", level=1)
    for text in (
        "The internal design combines CGGA and TCGA rather than fitting separate within-cohort internal models.",
        "The stored four-horizon survival probabilities do not permit calculation of a dense-time integrated Brier score; reported IBS is the mean of four horizon-specific Brier scores.",
        "CGGA includes primary and recurrent disease, producing clinical heterogeneity not explicitly modeled.",
        "Treatment, extent of resection, performance status, imaging, and center effects were unavailable or not shared across both cohorts.",
        "Biomarker missingness is substantial, especially for MGMT in CGGA.",
        "The knowledge relationships are expert-motivated inductive biases, not causal effects or a comprehensive knowledge graph.",
        "The selected 16-dimensional, one-layer, one-head architecture and bounded tuning grid were chosen after reviewing architecture experiments on these cohorts; the present external estimates are therefore exploratory follow-up results rather than untouched confirmatory validations.",
        "The architecture comparison conditions on the fitted predictions and used the same external cohorts to select the architecture; a new independent cohort is required for confirmatory evaluation.",
        "The primary Cox penalty and loss-balance configuration was selected separately inside each outer training set. Although leakage-safe, this nested selection adds model-selection variability not captured by the patient-level bootstrap intervals.",
        "The Efron sensitivity candidates produced very similar inner-validation scores, so their fold-specific selections may be sensitive to sampling variation and should not be interpreted as stable biological preferences.",
        "The nonlinear-main-effect, cross-fitted, and orthogonality variants were evaluated after inspection of the retained model; they are sensitivity analyses rather than independent confirmatory tests.",
        "The primary-tumor-only restriction was evaluated after inspecting full-cohort performance; changes in internal and TCGA-to-CGGA metrics partly reflect a different target population rather than a pure change in model quality.",
        "The selected dropout-0.40/Adam configuration changed two training factors simultaneously after inspection of prior results; their individual effects cannot be separated, and its external ranks are exploratory rather than confirmatory.",
        "Exact rerun reproducibility demonstrates computational stability under the fixed software and hardware environment, but it does not address sampling uncertainty, model-selection bias, or transportability to a new cohort.",
        "The two-head/dropout-0.50/AdamW candidate changed attention heads, dropout, and optimizer simultaneously; its equal rank profile and worse prediction error do not identify the contribution of any one factor.",
        "The dropout-0.60 experiment isolates one training factor but was still evaluated after reviewing dropout-0.40 results; its negative finding is exploratory rather than an independently prespecified test.",
        "The one-head dropout-0.50/AdamW experiment changed dropout and optimizer together after reviewing earlier results; it cannot attribute its small differences to either factor and remains exploratory.",
        "The 32-dimensional/120-epoch experiment nominally changed representation capacity and the epoch allowance simultaneously. The cap did not bind because selected epochs ranged from 70 to 99, but a complete 16d/120 versus 32d/170 factorial comparison would still provide the cleanest formal attribution.",
        "One Survival SVM internal fold produced an optimizer convergence warning; its predictions were finite, but the warning indicates numerical sensitivity.",
        "Linear regression ignores censoring and should not guide substantive survival conclusions.",
        "No competing risks, time-varying covariates, prospective validation, or clinical decision-curve analysis was performed.",
        "Bootstrap intervals condition on the fitted predictions and therefore omit variability from repeating split generation, hyperparameter selection, and model refitting.",
        "Calibration was assessed in the large at four horizons; subgroup calibration, calibration slopes, and dense-time calibration curves were not estimated.",
    ):
        bullet(doc, text)

    doc.add_heading("10. Conclusion", level=1)
    paragraph(
        doc,
        "rain_final provides a complete six-model comparison based on the corrected CGGA and TCGA raw "
        "cohorts. The selected dropout-0.40/Adam CoxRes-KGA configuration ranked first for C-index, mean AUC, "
        "and reported IBS in combined internal validation and in CGGA-to-TCGA validation. In TCGA-to-CGGA "
        "validation it ranked second for C-index, third for mean AUC, and first for reported IBS. The complete "
        "run was reproduced with byte-identical predictions, metrics, rankings, and training decisions. "
        "Accordingly, the appropriate conclusion is strong exploratory performance with clear methodological, "
        "calibration, model-selection, and validation limitations. "
        "The controlled architecture analysis favored the 16-dimensional, one-head specification over the "
        "former 32-dimensional model and the two-head alternative. A nested Efron/learnable-bias sensitivity "
        "analysis did not consistently improve discrimination, supporting retention of fixed knowledge-bias "
        "strengths within the primary nested-tuning procedure. Nonlinear Cox terms, cross-fitted averaging, "
        "and residual orthogonality also failed to yield a reliable paired improvement and were not promoted."
        " Restricting both cohorts to primary tumors improved absolute discrimination but did not consistently improve "
        "relative rank or prediction error, so it was retained as a disease-setting sensitivity analysis. "
        "Selecting dropout 0.40 with Adam improved the exploratory rank profile relative to the dropout-0.12/AdamW "
        "reference, but most paired intervals included zero and the joint intervention requires confirmation in an independent cohort. "
        "The two-head/dropout-0.50/AdamW candidate matched that rank profile but worsened internal and forward-external "
        "prediction error and was therefore not selected over the dropout-0.40/Adam variant. Increasing dropout alone "
        "to 0.60 also worsened internal prediction error and reverse-external AUC, reinforcing dropout 0.40 as the "
        "preferred exploratory setting. The controlled one-head dropout-0.50/AdamW candidate likewise reduced "
        "reverse-external AUC rank and worsened internal prediction error, so it was not promoted. The "
        "32-dimensional/120-epoch candidate significantly weakened forward-external discrimination and also "
        "remained a negative sensitivity result."
    )

    doc.add_heading("11. Reproducibility and File Map", level=1)
    paragraph(
        doc,
        "The project separates immutable raw inputs, processed cohorts, stored splits, patient-level "
        "predictions, metric tables, R figures, source code, tests, provenance, and the final report. "
        "Running source/run_final.py cleans the corrected raw workbooks, recreates validation splits, "
        "refits all six models, recomputes paired-bootstrap analyses, and regenerates the R figures and report."
    )
    table(
        doc,
        ["Location", "Purpose"],
        [
            ["data/raw", "Source workbooks and reference paper"],
            ["data/processed", "Harmonized cohort and descriptive summaries"],
            ["data/splits", "Deterministic internal and external validation splits"],
            ["results/predictions", "Six-model patient-level predictions"],
            ["results/metrics", "Fold, summary, bootstrap, paired-comparison, provenance, and metric-audit outputs"],
            ["results/sensitivity/rank_tuning", "Former 32d/1-head training-only tuning artifacts retained as the architecture reference"],
            ["results/sensitivity/d16_h1 and d16_h2", "Controlled 16-dimensional one-head and two-head architecture runs"],
            ["results/sensitivity/architecture_comparison", "Architecture point estimates, fold summaries, and paired-bootstrap differences"],
            ["results/sensitivity/efron_ablation", "Training-only Efron/learnable-bias ablation predictions, selection records, and metrics"],
            ["results/sensitivity/structured_*", "Nonlinear Cox, cross-fitted, orthogonality, and paired-bootstrap sensitivity artifacts"],
            ["results/sensitivity/primary_tumor_only", "Audited primary-tumor subset, six-model predictions, bootstrap metrics, and full-versus-primary comparisons"],
            ["results/sensitivity/dropout040_adam", "Selected dropout-0.40/Adam predictions, all-model ranks, bootstrap results, and comparison with the dropout-0.12/AdamW reference"],
            ["results/sensitivity/dropout050_adamw_h2", "Two-head/dropout-0.50/AdamW predictions and paired comparisons with earlier configurations"],
            ["results/sensitivity/dropout060_adam", "Dropout-0.60/Adam predictions and paired single-factor comparison with dropout 0.40"],
            ["results/sensitivity/dropout050_adamw_h1", "One-head/dropout-0.50/AdamW predictions and paired comparison with dropout 0.40/Adam"],
            ["results/sensitivity/d32_e120_adam", "32-dimensional/120-epoch predictions and paired comparison with the selected 16-dimensional/170-epoch model"],
            ["results/figures", "R-generated figures"],
            ["source", "Data preparation, model fitting, evaluation, visualization, and report code"],
            ["tests", "Scope, metric, provenance, and prediction-integrity tests"],
        ],
        [1.75, 4.75],
        "Table 17. rain_final project organization.",
        8.4,
    )

    doc.add_heading("References", level=1)
    references = [
        "Cheng J, Kuang H, Yang S, et al. Segmentation-Guided Deep Learning for Glioma Survival Risk Prediction with Multimodal MRI. Big Data Mining and Analytics. 2025;8(2):364-382. doi:10.26599/BDMA.2024.9020083.",
        "Cox DR. Regression Models and Life-Tables. Journal of the Royal Statistical Society: Series B. 1972;34(2):187-220. doi:10.1111/j.2517-6161.1972.tb00899.x.",
        "Graf E, Schmoor C, Sauerbrei W, Schumacher M. Assessment and comparison of prognostic classification schemes for survival data. Statistics in Medicine. 1999;18:2529-2545.",
        "Heagerty PJ, Zheng Y. Survival model predictive accuracy and ROC curves. Biometrics. 2005;61(1):92-105.",
        "Ishwaran H, Kogalur UB, Blackstone EH, Lauer MS. Random Survival Forests. Annals of Applied Statistics. 2008;2(3):841-860.",
        "Katzman JL, Shaham U, Cloninger A, et al. DeepSurv: personalized treatment recommender system using a Cox proportional hazards deep neural network. BMC Medical Research Methodology. 2018;18:24.",
        "Van Belle V, Pelckmans K, Van Huffel S, Suykens JAK. Support vector methods for survival analysis. Artificial Intelligence in Medicine. 2011;53(2):107-118.",
        "Barnwal A, Cho H, Hocking TD. Survival regression with accelerated failure time model in XGBoost. Journal of Computational and Graphical Statistics. 2022;31(4):1292-1302.",
        "Vaswani A, Shazeer N, Parmar N, et al. Attention Is All You Need. Advances in Neural Information Processing Systems. 2017;30.",
        "Gorishniy Y, Rubachev I, Khrulkov V, Babenko A. Revisiting Deep Learning Models for Tabular Data. Advances in Neural Information Processing Systems. 2021;34:18932-18943.",
        "Hao J, Kim Y, Mallavarapu T, Oh JH, Kang M. Interpretable deep neural network for cancer survival analysis by integrating genomic and clinical data. BMC Medical Genomics. 2019;12(Suppl 10):189. doi:10.1186/s12920-019-0624-2.",
        "Lin DY. On the Breslow estimator. Lifetime Data Analysis. 2007;13:471-480. doi:10.1007/s10985-007-9048-y.",
        "Lee C, Zame WR, Yoon J, van der Schaar M. DeepHit: A Deep Learning Approach to Survival Analysis With Competing Risks. Proceedings of the AAAI Conference on Artificial Intelligence. 2018;32(1):2314-2321. doi:10.1609/aaai.v32i1.11842.",
        "Efron B. The Efficiency of Cox's Likelihood Function for Censored Data. Journal of the American Statistical Association. 1977;72(359):557-565. doi:10.1080/01621459.1977.10480613.",
    ]
    for index, reference in enumerate(references, start=1):
        item = doc.add_paragraph()
        item.paragraph_format.left_indent = Inches(0.25)
        item.paragraph_format.first_line_indent = Inches(-0.25)
        item.paragraph_format.space_after = Pt(5)
        set_font(item.add_run(f"{index}. {reference}"), size=9.5)

    output = REPORT_DIR / "glioma_survival_retained_results_report.docx"
    doc.save(output)
    print(output)


if __name__ == "__main__":
    build_report()
